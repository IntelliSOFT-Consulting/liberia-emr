#!/usr/bin/env python3
"""Build the LE-22 Sync & EIP architecture review document (.docx).

The markdown under docs/architecture/ and docs/adr/ remains the source of truth for
engineering. This produces the review artefact for MOH ICT, who do not read the repository.

Usage:  python3 scripts/docs/build-le22-docx.py [output.docx]
"""
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = sys.argv[1] if len(sys.argv) > 1 else "docs/architecture/LE-22-sync-eip-architecture.docx"

ACCENT = RGBColor(0x1F, 0x4E, 0x79)
WARN = RGBColor(0xB0, 0x30, 0x00)

doc = Document()
for s in ("Normal",):
    doc.styles[s].font.name = "Calibri"
    doc.styles[s].font.size = Pt(10.5)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def para(text="", bold=False, italic=False, size=10.5, color=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold, r.italic = bold, italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullets(items, style="List Bullet"):
    for it in items:
        p = doc.add_paragraph(style=style)
        if isinstance(it, tuple):
            r = p.add_run(it[0]); r.bold = True
            p.add_run(" " + it[1])
        else:
            p.add_run(it)
        p.paragraph_format.space_after = Pt(3)


def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ""
        r = c.paragraphs[0].add_run(hd)
        r.bold = True
        r.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def callout(label, text):
    p = doc.add_paragraph()
    r = p.add_run(f"{label}  ")
    r.bold = True; r.font.color.rgb = WARN; r.font.size = Pt(10)
    r2 = p.add_run(text); r2.font.size = Pt(10)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(8)


# ─────────────────────────── Title ───────────────────────────
t = doc.add_heading("LiberiaEMR Sync & EIP Architecture", level=0)
for r in t.runs:
    r.font.color.rgb = ACCENT
para("Facility → central synchronisation, patient identity and cross-facility query",
     italic=True, size=12)
para()
table(["", ""], [
    ["Ticket", "LE-22, Sync & EIP architecture design [CUSTOM DEV]"],
    ["Owner", "Paul Ekirapa, IntelliSOFT Consulting"],
    ["Reviewers", "IntelliSOFT engineering; MOH ICT Unit"],
    ["Status", "Draft for review"],
    ["Date", "18 August 2026"],
    ["Sign-off target", "21 August 2026"],
    ["Implements", "Sprint 3 (outbound push), Sprint 4 (cross-facility query)"],
], widths=[1.4, 4.8])

para()
para("This document covers the three design decisions the MOH must agree, the technology "
     "selection behind them, and the engineering risks that remain ours to close. It is the "
     "scaling template for every facility after Careysburg and Barnersville.", size=10.5)

doc.add_page_break()

# ─────────────────────── Executive summary ───────────────────────
h("1. Executive summary", 1)
para("Each facility runs a complete EMR and keeps working with no connectivity. Changes are "
     "captured from the database and pushed to a central instance, queueing locally when the "
     "link is down and draining automatically when it returns. Sync is one-way: central never "
     "writes back into a facility. Cross-facility query, in Sprint 4, is a read path, not a "
     "second write direction.")

h("1.1 What we are asking the MOH to decide", 2)
table(["#", "Decision", "Recommendation"], [
    ["1", "Patient identity across facilities",
     "Central assigns its own identifier and LINKS records. Never auto-merge."],
    ["2", "How much of a record crosses facilities",
     "Demographics plus a defined clinical summary, not full history."],
    ["3", "Offline registration and duplicate handling",
     "Registration never waits for central; duplicates resolved at central by a named MOH role."],
], widths=[0.35, 2.2, 3.65])

h("1.2 What we found that changes the plan", 2)
bullets([
    ("Nothing is built yet.",
     "The sync services are declared in configuration but no images are produced. This is good news; corrections are cheap now."),
    ("The technology choice is settled and mature.",
     "openmrs-dbsync already provides 34 mapped clinical entities, retry and conflict queues, payload encryption and metrics."),
    ("Two compatibility gaps must be tested before build.",
     "Our database (MariaDB) and platform version sit outside the module's documented support envelope."),
    ("One failure mode would silently lose data.",
     "If a facility publishes before the central receiver has subscribed, those messages are lost with no error. This is a start-order rule."),
])

h("1.3 The principle everything else follows", 2)
callout("PRINCIPLE",
        "Security fails closed. Clinical availability fails open. If security cannot be "
        "established the sync layer stops and queues. It never downgrades. If central is "
        "unreachable, the facility keeps delivering care, unaffected.")

doc.add_page_break()

# ─────────────────────── Topology ───────────────────────
h("2. Topology and how data moves", 1)
para("Facility instances are self-sufficient. The only dependency on central is the push "
     "itself, and that dependency is asynchronous.")

h("2.1 The path a change takes", 2)
bullets([
    ("1. A clinician records something.", "It is written to the facility database. Nothing else happens in that request; sync never sits in the clinical write path."),
    ("2. The change is captured from the database log.", "Not by polling and not through the application, so corrections applied outside the EMR are captured too."),
    ("3. It is queued locally.", "Durably. It survives a container restart, a power cut and a multi-week outage."),
    ("4. It is encrypted and published to central", "over a mutually authenticated TLS connection, when a link exists."),
    ("5. Central applies it", "to its aggregate database, keyed on the record's unique identifier so a retry converges rather than duplicating."),
])

h("2.2 When synchronisation happens", 2)
para("Continuously. There is no nightly batch and no 'sync now' button. When the link is up, "
     "a change reaches central within seconds. When it is down, changes accumulate in order and "
     "drain on reconnection with no action at the facility.")
bullets([
    ("Nobody has to remember to sync.", "A manual or scheduled sync fails quietly in the week a facility is busiest, which is when the data matters most."),
    ("Recovery needs no site visit.", "A clinic offline for three weeks reconnects and drains by itself. Essential for facilities a day's travel away."),
])

h("2.3 Components to be added", 2)
para("Three pieces exist in the design but not yet in the deployment configuration:")
table(["Component", "Purpose"], [
    ["Message broker (ActiveMQ Artemis)", "The transport the sync module requires, at central."],
    ["Sender management database", "Holds each facility's retry state and its position in the database log."],
    ["Facility client certificates", "The facility half of mutual TLS. Central already has its half."],
], widths=[2.1, 4.1])

doc.add_page_break()

# ─────────────────────── Technology ───────────────────────
h("3. Technology selection", 1)
para("Seven approaches were evaluated. The recommendation is openmrs-eip 4.2.0 with "
     "openmrs-dbsync 4.0.0 over ActiveMQ Artemis.")

h("3.1 Options considered", 2)
table(["Option", "Verdict", "Why"], [
    ["openmrs-dbsync", "RECOMMENDED", "Mature; most of the requirement already built and in production elsewhere."],
    ["Custom FHIR push", "Rejected", "Rebuilds a mature product. Programme enrolment has no honest FHIR mapping."],
    ["Legacy sync module", "Rejected", "Built for the old platform and UI; not a candidate."],
    ["sync2 module", "Rejected", "Never reached maturity; unmaintained."],
    ["Atom Feed + consumer", "Rejected", "Only sees changes made through the application; no delivery guarantee."],
    ["Database replication", "Rejected", "Copies credentials and local settings; wrong trust model."],
    ["Debezium + Kafka, hand-built", "Rejected for v1", "All the cost, none of the clinical domain logic."],
], widths=[1.5, 1.0, 3.7])

h("3.2 What we get without building it", 2)
para("Verified in the module's source code, not taken from its documentation:")
bullets([
    "34 OpenMRS clinical entities already modelled and mapped",
    "Durable retry queues on both sender and receiver",
    "A conflict queue that detects when central holds a newer version of a record",
    "Per-record hashes, the basis for proving facility and central actually match",
    "Payload encryption independent of the transport",
    "Prometheus metrics for operational monitoring",
    "Search-index maintenance and queue housekeeping at central",
])

h("3.3 Production precedent", 2)
para("This stack is not a paper choice. Mozambique's national EMR programme (CSaude) runs "
     "OpenMRS EIP-based synchronisation across many low-connectivity facilities, on an "
     "actively maintained deployment. The module's upstream is also active, with fixes "
     "landing as recently as late 2025.")

h("3.4 Identity: build now, adopt a registry later", 2)
para("Mature master patient index products exist (OpenCR, SanteMPI, JeMPI). We recommend "
     "building a small identity layer at central for this release, because deploying a "
     "national-scale registry to serve two facilities, before the matching policy is agreed, "
     "adds an operational burden without reducing risk.")
para("The identity layer is kept behind a clear boundary: central answers 'resolve these "
     "details to a person'. If the MOH later adopts a national registry, that boundary is "
     "repointed at it. The option is preserved rather than spent.", italic=True)
callout("NOTE",
        "One existing OpenMRS module does almost exactly what we need, but assumes every "
        "patient has a National ID. Liberia's is optional at registration, which is "
        "what makes our problem hard. It does not fit as-is.")

doc.add_page_break()

# ─────────────────────── What syncs ───────────────────────
h("4. What synchronises, and in what order", 1)
h("4.1 Coverage", 2)
para("All five planned data flows are covered by existing support. No custom development is "
     "needed to move patients, visits, encounters, observations or programme enrolments.")
table(["Data", "Covered", "Note"], [
    ["Patients, names, addresses, identifiers, relationships", "Yes", ""],
    ["Visits", "Yes", ""],
    ["Encounters, observations, diagnoses, providers", "Yes", ""],
    ["Conditions and allergies", "Yes", ""],
    ["MCH programme enrolments and states", "Yes", "Core Sprint 2/3 scope"],
    ["Orders (lab, drug, referral)", "At risk", "A known defect in the module affects order subtypes; to be verified"],
    ["Clinical metadata (concepts, locations)", "Not synced", "Delivered identically to both sides by the build; see 4.3"],
    ["User accounts", "References only", "Never passwords or security answers"],
], widths=[2.5, 0.9, 2.8])

h("4.2 Order", 2)
para("Changes are applied at central in the order the facility recorded them, per patient. A "
     "visit cannot land before its patient. Ordering across different patients is deliberately "
     "not enforced, because it would serialise the entire national push behind the slowest record.")
para("Where a change arrives before something it depends on (after a partial drain, for "
     "example) central holds it aside and retries when the dependency lands, rather than "
     "rejecting it or blocking everything behind it.")

h("4.3 Clinical metadata is not synchronised", 2)
para("Concepts, locations and forms are not pushed. Facility and central run the same built "
     "image, so they hold identical metadata with identical internal identifiers by "
     "construction. This is stronger than synchronising it, because it cannot drift at runtime.")
callout("RULE",
        "Facility and central must never run different content-package versions. A concept "
        "that exists at a facility but not at central is a sync failure. This belongs in the "
        "deployment runbook and the upgrade rehearsal.")

doc.add_page_break()

# ─────────────────────── Identity ───────────────────────
h("5. Decision 1: Patient identity across facilities", 1)
h("5.1 The problem, in one table", 2)
para("This is a property of MOH registration policy, not of the EMR:")
table(["Identifier", "Mandatory?", "Unique nationally?", "Consequence"], [
    ["MOH Health Record Number", "Yes", "No; issued per facility",
     "Both facilities will issue the same numbers to different people"],
    ["Liberia National ID", "No", "Yes, when present", "Often simply absent"],
    ["ANC Number", "No", "No; per facility", "Programme register number only"],
], widths=[1.6, 0.9, 1.4, 2.3])
para("So the mandatory identifier is not unique, and the unique identifier is not mandatory. "
     "There is no existing key that identifies a person nationally.", bold=True)

h("5.2 Recommendation: link, never merge", 2)
bullets([
    ("Central assigns a Central Person Identifier (CPI).", "Generated centrally on first receipt. Never typed by a clinician. Not printed on a card in this release."),
    ("Records are linked, not merged.", "Facility records are stored exactly as sent, with a separate table mapping them to a CPI."),
    ("Three outcomes, not two.", "High confidence links automatically; the middle band goes to a human; low confidence creates a new person."),
    ("National ID matches are checked, not trusted.", "A matching ID where sex disagrees is a transcription error, and goes to review."),
    ("Never match on name and date of birth alone.", "In maternal health this is not theoretical: twins share surname, date of birth, sex, mother's name, phone and facility."),
    ("Estimated dates of birth carry less weight.", "Many are a clerk's best estimate; treating them as fact manufactures confidence."),
])
callout("WHY THIS MATTERS CLINICALLY",
        "Merging two patients wrongly attaches one woman's obstetric history to another, a "
        "patient-safety event. With linking, a mistake is one row to correct. A merge destroys "
        "the evidence needed to detect it.")

h("5.3 How the identifier is generated and attached", 2)
bullets([
    ("It carries no meaning.", "No facility code, no birth year, no sex. An identifier encoding the facility discloses where someone sought care, and becomes wrong when they move."),
    ("It has a check character,", "so a mistyped identifier is rejected rather than creating a phantom person."),
    ("Every patient gets one immediately,", "before matching finishes; no record waits unidentified in a queue."),
    ("Identifiers are never reused or deleted.", "When two records are linked, one identifier becomes primary and the other permanently resolves to it. This is what makes the link reversible."),
    ("It is held at central, separately from the clinical copy,", "and attached when records are read, not written back into facility databases."),
])

h("5.4 What we need from MOH ICT", 2)
bullets([
    "Is a central identifier acceptable, or will National ID become mandatory at registration?",
    "Which named MOH role owns the duplicate review queue, and with what expected turnaround? "
    "Facility staff cannot resolve a duplicate whose other half is at another facility.",
])

doc.add_page_break()

# ─────────────────────── Scope ───────────────────────
h("6. Decision 2: How much of a record crosses facilities", 1)
para("A legal and proportionality question under the draft Data Protection Act, as much as a "
     "clinical one.")
table(["Option", "Clinical value", "Exposure"], [
    ["A: Demographics only", "Low; changes no decision", "Minimal"],
    ["B: Demographics + defined summary (RECOMMENDED)", "High; covers the first five minutes", "Bounded, and written down"],
    ["C: Full clinical history", "Marginally more than B", "Everything, to every facility"],
], widths=[2.4, 2.0, 1.8])

h("6.1 Recommended scope", 2)
bullets([
    "Active problems and conditions", "Allergies and intolerances", "Current medications",
    "Immunisations", "MCH programme enrolments and current state",
    "Last antenatal contact summary: gestational age, risk flags, next appointment",
    "A list of encounter dates, types and facilities, without the observations inside them",
])
para("The deciding argument is not volume. An enumerated list can be tested automatically; "
     "'clinical history' cannot be checked against anything, which makes it unenforceable "
     "however it is written in a policy.", bold=True)

h("6.2 Conditions attached", 2)
bullets([
    ("Query, never copy.", "The remote record is displayed, never stored locally."),
    ("One patient at a time.", "No export, no lists, no bulk retrieval."),
    ("A reason for access is captured", "and stored with the audit record."),
    ("Every access is logged,", "including those returning nothing, readable only by the ICT audit role."),
])

h("6.3 What we need from MOH ICT", 2)
bullets([
    "Confirmation of the list above.",
    "Any category to be excluded from cross-facility visibility (HIV status being the obvious candidate); this must be named before build.",
    "The lawful basis, and whether consent is captured at query time or covered by the care relationship.",
])

doc.add_page_break()

# ─────────────────────── Offline ───────────────────────
h("7. Decision 3: Offline registration, and the offline guarantee", 1)
h("7.1 Registration never waits", 2)
para("A clerk registers a patient with no link to central. The identifier assigned by central "
     "arrives later, asynchronously. No registration, search or clinical workflow blocks on it; "
     "an EMR that needs the network to register a patient is not offline-first.")

h("7.2 Two different duplicate problems", 2)
table(["", "Within one facility", "Across facilities"], [
    ["Visible to", "Facility staff", "Central only"],
    ["Cause", "The search missed an existing record", "No shared identifier at registration"],
    ["Resolved by", "Facility staff, locally", "A named MOH role, at central"],
], widths=[1.1, 2.4, 2.7])

h("7.3 The guarantee, and what could break it", 2)
para("The requirement is blunt: a facility with poor connectivity must lose nothing and "
     "reconcile completely once the link is stable, however long the gap. Retry logic alone "
     "does not deliver that. Each failure below has a specific closure.")
table(["Failure", "Trigger", "Closure"], [
    ["Database change log pruned before it is sent", "Outage longer than the log retention",
     "Six-month retention floor, with alerting on the margin"],
    ["Facility disk fills, and the database stops", "Long outage on an undersized disk",
     "Size for the full window; separate volume; alarms at 60/75/85%"],
    ["Certificate expires during the outage", "A long outage crossing an expiry date",
     "Long-lived certificates; alerts at 90/60/30 days"],
    ["Old data overwrites newer data", "Delayed or replayed delivery",
     "Central rejects updates older than what it holds"],
    ["One bad message blocks everything behind it", "A malformed record",
     "Limited retries, then set aside for a human; the stream continues"],
    ["A facility goes quiet and nobody notices", "Facility down, or sender crashed",
     "Heartbeat per facility, with alerting on silence"],
    ["Everything retried, yet records still missing", "Any of the above, or a defect",
     "Scheduled reconciliation comparing facility and central directly"],
    ["Facilities all reconnect at once", "A regional outage ends",
     "Staggered reconnection and per-facility rate limits"],
], widths=[1.8, 1.7, 2.7])

callout("THE TWO THAT GET DEFERRED",
        "Disk sizing is a clinical-safety requirement, not capacity planning; it is the only "
        "path where the sync layer can stop care. And retries prove only that what entered the "
        "pipeline left it; only reconciliation proves nothing was lost.")

h("7.4 How we will prove it", 2)
para("Before go-live, and repeated for every new facility. Tested, not argued:")
bullets([
    "Bring a facility up, sync, confirm it matches central.",
    "Sever the link and record a realistic clinical day.",
    "Keep it severed well past the longest outage the MOH expects to tolerate, watching disk.",
    "Restore the link with no intervention at the facility. A facility needing an engineer to resume does not meet the requirement.",
    "Confirm: everything drains, nothing diverges, no duplicates, and no patient data appears in any log.",
    "Repeat with two facilities reconnecting simultaneously.",
])

doc.add_page_break()

# ─────────────────────── Security ───────────────────────
h("8. Security", 1)
para("The threat model is set by the deployment: facility servers sit in health centres, not "
     "a data centre. They are physically reachable, unattended overnight, and hold a complete "
     "copy of the facility's clinical record.")

h("8.1 Controls", 2)
table(["Control", "How it is met"], [
    ["Encryption in transit", "TLS 1.2+ on every facility-to-central connection"],
    ["Facility authentication", "Mutual TLS; each facility has its own certificate. Nothing is shared between facilities"],
    ["Message-level encryption", "Payloads encrypted so that only central can read them, protecting queued and backed-up data, not just the connection"],
    ["Revocation", "Enforced at central, so a stolen facility server can be cut off without reaching it"],
    ["Isolation between facilities", "Each facility may only send, and only on its own channel. It cannot read another facility's data"],
    ["Audit", "Sync outcomes and every cross-facility access, readable only by the ICT audit role"],
    ["No patient data in logs", "Identifiers and outcomes only, including in error logs, where leaks usually occur"],
    ["Data at rest", "Full-disk encryption on facility servers; encrypted backups covering the queue and change log as well as the database"],
], widths=[1.9, 4.3])

h("8.2 Two points that are easy to miss", 2)
bullets([
    ("Facilities share one broker, so permissions are a national-scale control.",
     "A single misconfigured permission would let one facility read another's clinical data. It fails silently, so it is verified by testing that access is refused, not by reviewing the configuration."),
    ("Patient data exists at rest in more places than the database.",
     "The database change log holds months of every change; the retry queue and the broker hold clinical messages. All must be encrypted and backed up together."),
])

h("8.3 New controls for the register", 2)
para("Three controls arise from this design and are not yet in the MOH ICT control mapping: "
     "per-facility broker authorisation, full-disk encryption on facility servers, and "
     "certificate revocation enforced at central. Each needs an owner.")

doc.add_page_break()

# ─────────────────────── Risks & plan ───────────────────────
h("9. Risks and plan", 1)
h("9.1 Engineering risks: ours to close, not MOH decisions", 2)
table(["Risk", "Severity", "How it is closed"], [
    ["Our database may not be supported by the module's change-capture component",
     "Highest", "A test, before anything is built. If it fails, we change database platform now; nearly free today, a data migration after go-live"],
    ["Facility disk filling during a long outage would stop the database",
     "Highest", "Disk sizing in the hardware specification; separate volume; alarms"],
    ["Broker permissions could expose one facility's data to another",
     "Highest", "Send-only per facility, proven by a test that access is refused"],
    ["Messages published before central is listening are lost silently",
     "Highest", "Durable subscription plus an enforced start order"],
    ["Our platform version is newer than the module documents",
     "High", "A schema comparison across the affected tables"],
    ["Order records have a known defect in the module",
     "High", "Verify; defer orders from the first release if it stands"],
    ["Review queues with no owner", "High", "A named MOH role with an expected turnaround"],
    ["No plan for loading a facility's pre-existing records", "High",
     "A one-time load during onboarding, one facility at a time, verified by comparing facility and central directly"],
], widths=[2.4, 0.8, 3.0])

h("9.2 Sequence", 2)
table(["Step", "Work", "Outcome"], [
    ["0", "Database compatibility test", "Go / no-go on the current database platform"],
    ["1", "Schema comparison for our platform version", "Confirms no gaps, or a known list"],
    ["2", "Pin versions; build the sync images", "Deployable artefacts; none exist today"],
    ["3", "Add broker and management database", "Stack runs end to end"],
    ["4", "Certificates and encryption keys", "Secure channel proven"],
    ["5", "Confirm data coverage", "Final list of what syncs"],
    ["6", "Initial load of each facility's existing records", "Historical data at central, verified by reconciliation"],
    ["7", "Identity layer at central", "Duplicates surfaced, never merged"],
    ["8", "Offline acceptance test", "The guarantee proven rather than claimed"],
    ["9", "Cross-facility query (Sprint 4)", "Read path, scoped as agreed"],
], widths=[0.5, 2.6, 3.1])

h("9.3 Decisions required by 21 August", 2)
table(["#", "Question", "Blocks"], [
    ["1", "Central identifier accepted, or National ID made mandatory?", "Identity design, Sprint 3"],
    ["2", "Who owns the duplicate review queue?", "Identity design, go-live"],
    ["3", "Is the pulled-record scope confirmed?", "Cross-facility query, Sprint 4"],
    ["4", "Any category excluded from cross-facility visibility?", "Scope, clinical content"],
    ["5", "Lawful basis, and consent at query time?", "Scope"],
    ["6", "Who owns the certificate and key lifecycle?", "Sprint 3"],
    ["7", "Longest facility outage we must tolerate?", "Disk sizing, log retention, hardware"],
    ["8", "Audit retention for cross-facility access?", "Audit control"],
    ["9", "Full-disk encryption on facility servers: accepted, and owned by whom?", "Control register"],
], widths=[0.35, 3.5, 2.35])
callout("QUESTION 7 SETS THE MOST",
        "It fixes log retention, facility disk size and the acceptance test. If the answer is "
        "unknown, treat it as the longest plausible outage; under-sizing causes silent data "
        "loss or a stopped database, while over-sizing costs only disk.")

doc.add_page_break()
h("Appendix: Source material", 1)
para("Engineering detail is maintained in the LiberiaEMR repository and remains the source of "
     "truth:")
bullets([
    "docs/architecture/sync-eip.md, the full architecture",
    "docs/architecture/sync-module-evaluation.md, options, justification, plan and risks",
    "docs/architecture/sync-entity-coverage.md, what syncs, in what order, and custom work",
    "docs/adr/0005 (identity reconciliation), 0007 (pulled-record scope), 0008 (module selection)",
])
para()
para("External references, retrieved 18 August 2026:", bold=True)
bullets([
    "openmrs/openmrs-eip and mekomsolutions/openmrs-dbsync, source and documentation",
    "Debezium documentation, connector support for MariaDB and MySQL",
    "OpenHIE Client Registry specification; OpenCR, SanteMPI, JeMPI",
    "csaude/openmrs-module-mpi, an OpenMRS master patient index integration",
])

doc.save(OUT)
print(f"Wrote {OUT}")
